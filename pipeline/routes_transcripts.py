"""Move 2d (D-70) — deposition digest: map-reduce over ALL of one transcript's pages,
where only MECHANICALLY VERIFIED bullets survive.

POST /transcripts/{doc_id}/digest (SSE): walk every chunk (= page) of one transcript in
page order, in batches of ~10 pages per local LLM call. Each batch prompt asks for
topic-tagged bullets, each carrying a verbatim quote in the standard §10 citation tag.
EVERY bullet is then individually run through the UNCHANGED verifier against that
batch's own grounding — an unverified bullet is dropped and counted, never displayed.
The reduce step only GROUPS and ORDERS verified bullets by topic; it generates no new
text (the silent-partiality failure of top-k "summaries" is structurally excluded:
this reads all N pages and says so).

POST /transcripts/digest.docx: render a digest JSON to a Word table (Topic | Verbatim
testimony | Cite) — the artifact that goes in a trial binder, not a chat bubble.

Single-document, matter comes from the catalog row; loopback LLM only.
"""

import io
import json
import re

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse, Response

import activity
import catalog
import routes_kb
import transcript_extract as te
from answering import _build_messages, _chat, _THINK_RE
from embed_store import open_table
from verifier import verify_answer

router = APIRouter()

PAGES_PER_BATCH = 10
_BULLET_RE = re.compile(r"^[-*]\s*\[(?P<topic>[^\]]{2,60})\]\s*(?P<body>.+)$")

_DIGEST_INSTRUCTIONS = (
    "From the transcript pages above, extract the key testimony as bullets. Format "
    "EVERY bullet exactly as:\n"
    "- [topic] short statement with the verbatim quote [document: <filename>, page: "
    "<page_number>, chunk: <chunk_id>, span: \"<verbatim quote>\"]\n"
    "Rules: one bullet per distinct fact; topic is 1-3 words (e.g. [maintenance fees], "
    "[timeline], [safety]); the span must be QUOTED VERBATIM from the pages above; "
    "skip procedural chatter (scheduling, objections without substance). If nothing "
    "substantive appears, output exactly: NONE"
)


def _transcript_chunks(doc):
    """All chunks of one transcript document, in page order, from the KB store."""
    table = open_table(str(routes_kb.KB_DB))
    fn = doc["filename"].replace("'", "''")
    mt = doc["matter_slug"].replace("'", "''")
    n = table.count_rows()
    rows = (table.search()
            .select(["source_filename", "matter", "page_number", "section",
                     "char_start", "char_end", "text"])
            .where(f"source_filename = '{fn}' AND matter = '{mt}'", prefilter=True)
            .limit(max(n, 1)).to_arrow().to_pylist())
    return sorted(rows, key=lambda r: r["page_number"])


def _grounding(rows):
    return [{"chunk_id": f"C{i}", "source_filename": r["source_filename"],
             "page_number": r["page_number"], "char_start": r["char_start"],
             "char_end": r["char_end"], "text": r["text"]}
            for i, r in enumerate(rows, 1)]


def _digest_batch(rows, doc):
    """One batch: generate bullets, verify EACH ONE mechanically, derive page:line.
    Returns (verified_bullets, rejected_count)."""
    grounding = _grounding(rows)
    chunks_for_prompt = [{**g, "matter": doc["matter_slug"], "section": ""} for g in
                         [{**r, "source_filename": r["source_filename"]} for r in rows]]
    messages, _ = _build_messages(chunks_for_prompt, _DIGEST_INSTRUCTIONS)
    raw = _THINK_RE.sub("", _chat(messages)).strip()
    verified, rejected = [], 0
    for line in raw.splitlines():
        line = line.strip()
        m = _BULLET_RE.match(line)
        if not m:
            continue
        verdict = verify_answer(line, grounding)
        if not verdict["citations"]:
            rejected += 1
            continue
        c = verdict["citations"][0]
        entries = catalog.line_map_for_page(doc["id"], c["page"])
        page_text = next((g["text"] for g in grounding
                          if g["page_number"] == c["page"]), "")
        lines = te.derive_lines(entries, c.get("char_start", -1), c.get("char_end", -1),
                                page_text, c.get("span", "")) if entries else None
        body = _BULLET_RE.match(line).group("body")
        body = re.sub(r"\[document:[^\]]*\]", "", body).strip()
        verified.append({"topic": m.group("topic").strip().casefold(),
                         "text": body, "span": c.get("span", ""),
                         "filename": c["filename"], "page": c["page"],
                         "lines": lines})
    return verified, rejected


@router.post("/transcripts/{doc_id}/digest")
def digest(doc_id: int):
    doc = catalog.get_document(doc_id)
    if not doc or doc.get("doc_type") != "transcript":
        raise HTTPException(status_code=404, detail="not a transcript document")
    if doc["status"] not in ("ready", "needs_review"):
        raise HTTPException(status_code=400, detail=f"document is {doc['status']}")
    rows = _transcript_chunks(doc)
    if not rows:
        raise HTTPException(status_code=400, detail="no indexed pages for this document")

    def event(name, obj):
        return f"event: {name}\ndata: {json.dumps(obj)}\n\n"

    def gen():
        batches = [rows[i:i + PAGES_PER_BATCH] for i in range(0, len(rows), PAGES_PER_BATCH)]
        yield event("meta", {"pages": len(rows), "batches": len(batches),
                             "filename": doc["filename"]})
        all_bullets, rejected_total = [], 0
        for bi, batch in enumerate(batches, 1):
            activity.mark_chat()   # digest is user-initiated foreground work
            try:
                bullets, rejected = _digest_batch(batch, doc)
            except Exception as e:  # one bad batch never kills the digest — count it
                yield event("batch", {"batch": bi, "of": len(batches),
                                      "error": f"{type(e).__name__}"})
                continue
            all_bullets.extend(bullets)
            rejected_total += rejected
            yield event("batch", {"batch": bi, "of": len(batches),
                                  "verified": len(bullets), "rejected": rejected})
        # reduce: group + order only; no new text is generated here
        topics = {}
        for b in all_bullets:
            topics.setdefault(b["topic"], []).append(b)
        digest_obj = {
            "filename": doc["filename"],
            "coverage": f"built from all {len(rows)} indexed pages",
            "topics": [{"topic": t, "bullets": sorted(bs, key=lambda b: (b["page"],))}
                       for t, bs in sorted(topics.items())],
            "stats": {"pages": len(rows), "bullets_verified": len(all_bullets),
                      "bullets_rejected_unverified": rejected_total},
        }
        yield event("done", digest_obj)

    return StreamingResponse(gen(), media_type="text/event-stream")


@router.post("/transcripts/digest.docx")
def digest_docx(payload: dict):
    """Render a digest JSON (from the SSE 'done' event) to a Word table for the trial
    binder. Verbatim testimony + cite per row; nothing is added or rewritten."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(f"Deposition Digest - {payload.get('filename', '')}", level=1)
    doc.add_paragraph(payload.get("coverage", ""))
    doc.add_paragraph("Every quote below was mechanically verified against the "
                      "transcript text. Verify context before use; this is not legal "
                      "advice.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Topic", "Testimony (verbatim quote)", "Cite"
    for t in payload.get("topics", []):
        for b in t.get("bullets", []):
            row = table.add_row().cells
            row[0].text = t.get("topic", "")
            row[1].text = f"{b.get('text', '')}\n\"{b.get('span', '')}\""
            cite = f"{b.get('filename', '')} p.{b.get('page', '')}"
            if b.get("lines"):
                cite += f":{b['lines']}"
            row[2].text = cite
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    buf = io.BytesIO()
    doc.save(buf)
    return Response(
        buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition":
                 f'attachment; filename="digest-{payload.get("filename", "transcript")}.docx"'})
